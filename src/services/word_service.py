import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from typing import List, Tuple
from io import BytesIO
from src.services.log_service import logger

class WordService:
    """Word 문서 생성 서비스"""
    
    def __init__(self, template_file: str = "3677.docx"):
        self.template_file = template_file
        # 바코드 크기 (MM 단위로 통일)
        self.barcode_width_mm = 30.0  # 기본 30mm
        self.barcode_height_mm = 15.0  # 기본 15mm
        self.text_font_size = 0.08
        self.font_name = "맑은 고딕"
        self.highlight_color = RGBColor(255, 255, 0)
    
    def set_barcode_size_mm(self, width_mm: float, height_mm: float):
        """바코드 크기를 MM 단위로 설정"""
        self.barcode_width_mm = width_mm
        self.barcode_height_mm = height_mm
        logger.info("WordService", f"바코드 크기 설정: {width_mm:.1f}mm x {height_mm:.1f}mm")
    
    def _mm_to_inches(self, mm: float) -> float:
        """MM를 인치로 변환"""
        return mm / 25.4
    
    def create_label_page(self, items_for_page: List[Tuple[str, str, str, str]], 
                         page_name: str, barcode_images: dict, output_dir: str = "output") -> bool:
        """한 페이지 분량의 라벨을 생성하고 파일로 저장"""
        try:
            # 출력 디렉토리 생성
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            
            # 새 문서 생성 (템플릿 복사)
            page_doc = Document(self.template_file)
            
            # 첫 번째 테이블 찾기
            page_table = None
            for tbl in page_doc.tables:
                page_table = tbl
                break
            
            if page_table is None:
                logger.error("WordService", "템플릿에서 테이블을 찾을 수 없습니다!")
                return False
            
            # 테이블 채우기
            item_idx = 0
            for r in range(len(page_table.rows)):
                for c in range(len(page_table.columns)):
                    if item_idx < len(items_for_page):
                        name, price, category, code = items_for_page[item_idx]
                        
                        # 메모리에서 바코드 이미지 가져오기
                        if code in barcode_images:
                            cell = page_table.cell(r, c)
                            
                            # 기존 내용 지우기
                            cell.text = ""
                            
                            # 바코드 이미지 추가 (맨 위, 중앙 정렬)
                            p1 = cell.paragraphs[0]
                            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run1 = p1.add_run()
                            
                            # 바코드 이미지 가져오기 (파일 경로 또는 메모리 버퍼)
                            barcode_data = barcode_images[code]
                            
                            if isinstance(barcode_data, str):
                                # 파일 경로인 경우
                                run1.add_picture(barcode_data, 
                                               width=Inches(self._mm_to_inches(self.barcode_width_mm)), 
                                               height=Inches(self._mm_to_inches(self.barcode_height_mm)))
                            else:
                                # BytesIO 메모리 버퍼인 경우
                                barcode_data.seek(0)  # 버퍼 위치를 처음으로 리셋
                                run1.add_picture(barcode_data, 
                                               width=Inches(self._mm_to_inches(self.barcode_width_mm)), 
                                               height=Inches(self._mm_to_inches(self.barcode_height_mm)))
                            
                            # 상품명과 가격을 위한 새 문단 추가
                            p2 = cell.add_paragraph()
                            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # 상품명 부분
                            name_run = p2.add_run(name)
                            name_run.font.size = Inches(self.text_font_size)
                            name_run.font.name = self.font_name
                            
                            # 공백
                            space_run = p2.add_run(" ")
                            space_run.font.size = Inches(self.text_font_size)
                            space_run.font.name = self.font_name
                            
                            # 가격 부분 (글자 테두리)
                            price_run = p2.add_run(f"{price}₩")
                            price_run.font.size = Inches(self.text_font_size)
                            price_run.font.name = self.font_name
                            price_run.font.bold = True
                            
                            # 글자 테두리 설정
                            rPr = price_run._element.get_or_add_rPr()
                            bdr = OxmlElement('w:bdr')
                            bdr.set(qn('w:val'), 'single')
                            bdr.set(qn('w:sz'), '4')
                            bdr.set(qn('w:space'), '0')
                            bdr.set(qn('w:color'), '000000')
                            rPr.append(bdr)
                            
                            item_idx += 1
                        else:
                            logger.warning("WordService", f"바코드 이미지를 찾을 수 없음: {code}")
                            # 바코드 이미지가 없어도 텍스트 정보는 추가
                            cell = page_table.cell(r, c)
                            cell.text = ""
                            
                            # 상품명과 가격만 추가
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # 상품명 부분
                            name_run = p.add_run(name)
                            name_run.font.size = Inches(self.text_font_size)
                            name_run.font.name = self.font_name
                            
                            # 공백
                            space_run = p.add_run(" ")
                            space_run.font.size = Inches(self.text_font_size)
                            space_run.font.name = self.font_name
                            
                            # 가격 부분
                            price_run = p.add_run(f"{price}₩")
                            price_run.font.size = Inches(self.text_font_size)
                            price_run.font.name = self.font_name
                            price_run.font.bold = True
                            
                            # 바코드 번호도 표시
                            code_run = p.add_run(f"\n{code}")
                            code_run.font.size = Inches(self.text_font_size * 0.8)
                            code_run.font.name = self.font_name
                            
                            item_idx += 1
                    else:
                        # 빈 셀 처리
                        cell = page_table.cell(r, c)
                        cell.text = ""
            
            # 파일명에서 특수문자 제거
            safe_name = "".join(c for c in page_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = os.path.join(output_dir, f"{safe_name}_label.docx")
            
            # 파일 저장
            page_doc.save(filename)
            logger.info("WordService", f"페이지 저장 완료: {filename} ({len(items_for_page)}개 라벨)")
            return True
            
        except Exception as e:
            logger.error("WordService", f"라벨 페이지 생성 실패: {e}")
            return False

    def get_table_max_size(self, template: str):
        """
        Prints the number of rows and columns of the first table in a .docx file.
        """
        try:
            document = Document(template)
            if not document.tables:
                print("No tables found in the document.")
                return 0

            table = document.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            return num_cols * num_rows

        except Exception as e:
            print(f"An error occurred: {e}")
    
    def generate_label_documents(self, items: List[Tuple[str, str, str, str]], 
                                barcode_images: dict, output_dir: str = "output") -> int:
        """상품별로 그룹화하여 라벨 문서 생성"""
        if not items:
            return 0
        
        # 한 페이지당 라벨 수 (78개)
        labels_per_page = 78
        
        # 상품별로 그룹화하여 처리
        current_product = None
        current_items = []
        total_files_created = 0
        idx = 0
        
        while idx < len(items):
            name, price, category, code = items[idx]
            
            # 새로운 상품이 시작되거나 현재 상품의 라벨이 78개에 도달한 경우
            if current_product != name or len(current_items) >= labels_per_page:
                # 이전 상품의 라벨들이 있으면 파일로 저장
                if current_items:
                    if self.create_label_page(current_items, current_product, barcode_images, output_dir):
                        total_files_created += 1
                    current_items = []
                
                current_product = name
            
            # 현재 아이템을 리스트에 추가
            current_items.append((name, price, category, code))
            idx += 1
        
        # 마지막 상품 처리
        if current_items:
            if self.create_label_page(current_items, current_product, barcode_images, output_dir):
                total_files_created += 1
        
        print(f"\n=== 작업 완료 ===")
        print(f"총 {total_files_created}개 파일 생성됨")
        print(f"총 {len(items)}개 라벨 처리됨")
        
        return total_files_created
    
    def generate_single_label_document(self, items: List[Tuple[str, str, str, str]], 
                                     barcode_images: dict, output_dir: str = "output") -> int:
        """모든 라벨을 하나의 문서로 생성 (모든 상품의 라벨을 순서대로 배치)"""
        if not items:
            return 0
        
        print(f"통합 문서 생성 시작 - 총 {len(items)}개 라벨")
        
        # 첫 번째 페이지 생성
        pages_created = []
        current_items = []
        
        # 템플릿에서 한 페이지당 라벨 수 계산
        try:
            template_doc = Document(self.template_file)
            template_table = None
            for tbl in template_doc.tables:
                template_table = tbl
                break
            
            if template_table is None:
                print("템플릿에서 테이블을 찾을 수 없습니다!")
                return 0
            
            rows_per_page = len(template_table.rows)
            cols_per_page = len(template_table.columns)
            labels_per_page = rows_per_page * cols_per_page
            
            print(f"템플릿 정보: {rows_per_page}행 x {cols_per_page}열 = {labels_per_page}개 라벨/페이지")
            
        except Exception as e:
            print(f"템플릿 분석 실패: {e}")
            return 0
        
        # 라벨들을 페이지별로 나누어 처리
        page_num = 1
        for i in range(0, len(items), labels_per_page):
            page_items = items[i:i + labels_per_page]
            
            # 각 페이지를 개별 파일로 생성
            page_name = f"통합_라벨_페이지_{page_num}"
            if self.create_label_page(page_items, page_name, barcode_images, output_dir):
                pages_created.append(f"{page_name}_label.docx")
                print(f"페이지 {page_num} 생성 완료 ({len(page_items)}개 라벨)")
            
            page_num += 1
        
        # 생성된 페이지들을 하나의 문서로 합치기
        if pages_created:
            try:
                # 첫 번째 문서를 기본으로 사용
                first_page_path = os.path.join(output_dir, pages_created[0])
                combined_doc = Document(first_page_path)
                
                # 나머지 페이지들을 추가
                for page_file in pages_created[1:]:
                    page_path = os.path.join(output_dir, page_file)
                    page_doc = Document(page_path)
                    
                    # 페이지 나누기 추가
                    combined_doc.add_page_break()
                    
                    # 페이지의 모든 요소를 복사
                    for element in page_doc.element.body:
                        combined_doc.element.body.append(element)
                
                # 통합 문서 저장
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                combined_filename = os.path.join(output_dir, f"통합_라벨_{timestamp}.docx")
                combined_doc.save(combined_filename)
                
                # 임시 페이지 파일들 삭제
                for page_file in pages_created:
                    try:
                        os.remove(os.path.join(output_dir, page_file))
                    except:
                        pass
                
                print(f"통합 라벨 문서 저장 완료: {combined_filename}")
                print(f"총 {len(items)}개 라벨이 {len(pages_created)}페이지에 생성됨")
                
                return 1
                
            except Exception as e:
                print(f"문서 합치기 실패: {e}")
                # 실패 시 개별 페이지 파일들은 그대로 유지
                return len(pages_created)
        
        return 0

    def get_cell_size_mm(self, template: str) -> Tuple[float, float]:
        """
        첫 번째 테이블의 한 셀(cell) 너비와 높이를 mm 단위로 반환합니다.
        반환값: (width_mm, height_mm). 값을 찾지 못하면 0.0으로 반환합니다.

        구현 노트:
        - 테이블 열 너비는 <w:tblGrid>/<w:gridCol w:w="..."/> 값들을 사용해 계산합니다.
          이 값들은 일반적으로 "twips" (1/1440 인치) 단위로 표현되므로 mm로 변환합니다.
        - 행 높이는 각 <w:trPr>/<w:trHeight w:val="..."/> 값을 우선으로 시도합니다.
          없을 경우 모든 행의 trHeight 평균값을 시도합니다.
        - OOXML 문서에 따라 단위/속성 위치가 다를 수 있으므로 여러 속성명을 시도합니다.
        """
        try:
            document = Document(template)
            if not document.tables:
                return (0.0, 0.0)

            table = document.tables[0]

            # --- 가로 너비 계산 (tblGrid의 gridCol 사용) ---
            width_mm = 0.0
            try:
                grid = table._tbl.tblGrid
                # gridCol 요소들 가져오기
                grid_cols = grid.findall(qn('w:gridCol'))
                col_vals = []
                for gc in grid_cols:
                    # 여러 속성명 시도
                    val = None
                    for attr in (qn('w:w'), 'w', 'w:w', 'val'):
                        if gc.get(attr):
                            val = gc.get(attr)
                            break
                    if val:
                        try:
                            col_vals.append(int(val))
                        except Exception:
                            pass
                if col_vals:
                    # col_vals는 각 열의 폭(일반적으로 twips 또는 dxa) -> 평균 셀 너비 도출
                    # twip(1/1440 inch) 가정: mm = twip / 1440 * 25.4
                    avg_twips = sum(col_vals) / len(col_vals)
                    width_mm = avg_twips / 1440.0 * 25.4
            except Exception:
                width_mm = 0.0

            # --- 세로(행) 높이 계산 (첫 번째 행 또는 모든 행 trHeight 평균) ---
            height_mm = 0.0
            try:
                # 우선 첫 행의 trHeight 시도
                first_row = table.rows[0]
                trpr = first_row._tr.find(qn('w:trPr'))
                if trpr is not None:
                    trh = trpr.find(qn('w:trHeight'))
                    if trh is not None:
                        val = trh.get('val') or trh.get(qn('w:val')) or trh.get('w')
                        if val:
                            try:
                                height_mm = int(val) / 1440.0 * 25.4
                            except Exception:
                                height_mm = 0.0

                # 못 찾았으면 모든 행의 trHeight 평균 시도
                if not height_mm:
                    total_twips = 0
                    count = 0
                    for r in table.rows:
                        trpr = r._tr.find(qn('w:trPr'))
                        if trpr is None:
                            continue
                        trh = trpr.find(qn('w:trHeight'))
                        if trh is None:
                            continue
                        val = trh.get('val') or trh.get(qn('w:val')) or trh.get('w')
                        if val:
                            try:
                                total_twips += int(val)
                                count += 1
                            except Exception:
                                pass
                    if count > 0:
                        height_mm = (total_twips / count) / 1440.0 * 25.4
            except Exception:
                height_mm = 0.0

            return (round(width_mm, 2), round(height_mm, 2))

        except Exception as e:
            print(f"get_cell_size_mm error: {e}")
            return (0.0, 0.0)