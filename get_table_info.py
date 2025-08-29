import docx

def get_table_size(docx_path):
    """
    Prints the number of rows and columns of the first table in a .docx file.

    Args:
        docx_path (str): The path to the .docx file.
    """
    try:
        document = docx.Document(docx_path)
        if not document.tables:
            print("No tables found in the document.")
            return 0

        table = document.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        return num_cols * num_rows

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    get_table_size("templates/3677.docx")
