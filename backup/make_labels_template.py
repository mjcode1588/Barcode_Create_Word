import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from barcode import Code128
from barcode.writer import ImageWriter
from PIL import Image

# ===== 설정 =====
CSV_FILE = "items.xlsx"  # CSV 파일 (상품명, 바코드번호)
TEMPLATE_FILE = "3677.docx"  # 기존 양식 파일
OUTPUT_DOC = "labels_filled.docx"
BARCODE_DIR = "barcodes"

# 바코드 이미지 설정 (GS1-128 고화질)
BARCODE_WIDTH = 1.2   # 바코드 이미지 폭 (inch) - GS1-128에 맞게 조정
BARCODE_HEIGHT = 0.6  # 바코드 이미지 높이 (inch) - 고화질에 맞게 조정
TEXT_FONT_SIZE = 0.08  # 텍스트 폰트 크기 (inch)
FONT_NAME = "맑은 고딕"  # 폰트 이름 (예쁜 한글 폰트)
HIGHLIGHT_COLOR = RGBColor(255, 255, 0)  # 노란색 하이라이트
# =================

# 폴더 생성 (기존 바코드 파일들 삭제 후 재생성)
if os.path.exists(BARCODE_DIR):
    import shutil
    shutil.rmtree(BARCODE_DIR)
os.makedirs(BARCODE_DIR)

# Excel 파일 읽기
wb = load_workbook(CSV_FILE)
ws = wb.active

# 헤더 읽기 (첫 번째 행)
headers = []
for cell in ws[1]:
    headers.append(cell.value)
print(f"Excel 파일 컬럼: {headers}")

# 종류별 바코드 번호 생성
category_counters = {}
items = []
categories = []

# 데이터 읽기 (2번째 행부터)
for row in ws.iter_rows(min_row=2, values_only=True):
    if row[0] is None:  # 빈 행이면 건너뛰기
        continue
        
    name = str(row[0]).strip() if row[0] else ""
    price = str(row[1]).strip() if row[1] else ""
    category = str(row[2]).strip() if row[2] else ""
    copy_flag = row[3] if len(row) > 3 else False  # 복사 열 (4번째 열)
    
    # 종류별 카운터 관리
    if category not in category_counters:
        category_counters[category] = 1

    # 카테고리 인덱스 관리: categories 리스트에 없으면 추가
    if category not in categories:
        categories.append(category)

    category_index = categories.index(category) + 1  # 1-based index

    # 종류별 바코드 생성: PPON-{종류인덱스}{6자리순번} 형태 (전체 대문자)
    item_number = str(category_counters[category]).zfill(6)  # 6자리 순번
    barcode_number = f"PPON-{category_index}{item_number}".upper()
    
    # 복사 플래그에 따라 개수 결정
    if copy_flag:
        # True면 78개 추가
        for i in range(78):
            items.append((name, price, category, barcode_number))
    else:
        # False면 1개만 추가
        items.append((name, price, category, barcode_number))
    
    category_counters[category] += 1

print(f"총 {len(items)}개 상품 로드됨")
for category, count in category_counters.items():
    idx = categories.index(category) + 1 if category in categories else "?"
    print(f"  {category} (index:{idx}): {count-1}개")

# 바코드 이미지 생성 (GS1-128, 고화질)
unique_codes = set()
for name, price, category, code in items:
    if code not in unique_codes:
        unique_codes.add(code)
        filename = os.path.join(BARCODE_DIR, f"{code}.png")
        if not os.path.exists(filename):
            # GS1-128 형식으로 바코드 생성 (고화질 설정)
            writer = ImageWriter()
            writer.dpi = 300  # 고화질 DPI 설정
            writer.module_width = 0.2  # 바 너비 설정
            writer.module_height = 15.0  # 바 높이 설정
            writer.quiet_zone = 6.5  # 여백 설정
            writer.text_distance = 5.0  # 텍스트와 바코드 간격
            writer.font_size = 10  # 폰트 크기
            
            # GS1-128 바코드 생성 (Code128과 호환)
            barcode = Code128(code, writer=writer)
            # ImageWriter가 자동으로 .png를 추가하므로 확장자 없이 저장
            barcode.save(os.path.join(BARCODE_DIR, code))
            print(f"GS1-128 바코드 생성: {code} ({category} - {name})")

# 기존 템플릿 문서 열기
doc = Document(TEMPLATE_FILE)
print(f"템플릿 파일 로드: {TEMPLATE_FILE}")

# 첫 번째 테이블 찾기
table = None
for tbl in doc.tables:
    table = tbl
    break

if table is None:
    print("테이블을 찾을 수 없습니다!")
    exit(1)

print(f"테이블 크기: {len(table.rows)}행 x {len(table.columns)}열")

# 라벨 채우기 - 각 페이지를 별도 파일로 저장
idx = 0
labels_per_page = len(table.rows) * len(table.columns)  # 한 페이지당 라벨 수 (78개)

def create_label_page(items_for_page, page_name):
    """한 페이지 분량의 라벨을 생성하고 파일로 저장"""
    # 새 문서 생성 (템플릿 복사)
    page_doc = Document(TEMPLATE_FILE)
    
    # 첫 번째 테이블 찾기
    page_table = None
    for tbl in page_doc.tables:
        page_table = tbl
        break
    
    if page_table is None:
        print(f"템플릿에서 테이블을 찾을 수 없습니다!")
        return False
    
    # 테이블 채우기
    item_idx = 0
    for r in range(len(page_table.rows)):
        for c in range(len(page_table.columns)):
            if item_idx < len(items_for_page):
                name, price, category, code = items_for_page[item_idx]
                
                # 실제 존재하는 파일 경로 찾기
                img_path = os.path.join(BARCODE_DIR, f"{code}.png")
                if not os.path.exists(img_path):
                    img_path = os.path.join(BARCODE_DIR, f"{code}.png.png")
                
                if os.path.exists(img_path):
                    cell = page_table.cell(r, c)
                    
                    # 기존 내용 지우기
                    cell.text = ""
                    
                    # 바코드 이미지 추가 (맨 위, 중앙 정렬)
                    p1 = cell.paragraphs[0]
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run1 = p1.add_run()
                    run1.add_picture(img_path, width=Inches(BARCODE_WIDTH), height=Inches(BARCODE_HEIGHT))
                    
                    # 상품명과 가격을 위한 새 문단 추가
                    p2 = cell.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 중앙 정렬
                    
                    # 상품명 부분
                    name_run = p2.add_run(name)
                    name_run.font.size = Inches(TEXT_FONT_SIZE)
                    name_run.font.name = FONT_NAME
                    
                    # 공백
                    space_run = p2.add_run(" ")
                    space_run.font.size = Inches(TEXT_FONT_SIZE)
                    space_run.font.name = FONT_NAME
                    
                    # 가격 부분 (글자 테두리)
                    price_run = p2.add_run(f"{price}₩")
                    price_run.font.size = Inches(TEXT_FONT_SIZE)
                    price_run.font.name = FONT_NAME
                    price_run.font.bold = True  # 가격을 굵게 표시
                    
                    # 글자 테두리 설정
                    rPr = price_run._element.get_or_add_rPr()
                    bdr = OxmlElement('w:bdr')
                    bdr.set(qn('w:val'), 'single')
                    bdr.set(qn('w:sz'), '4')  # 테두리 두께
                    bdr.set(qn('w:space'), '0')
                    bdr.set(qn('w:color'), '000000')  # 검은색 테두리
                    rPr.append(bdr)
                    
                    item_idx += 1
                else:
                    print(f"바코드 이미지를 찾을 수 없음: {img_path}")
                    item_idx += 1
            else:
                # 빈 셀 처리
                cell = page_table.cell(r, c)
                cell.text = ""
    
    # 파일명에서 특수문자 제거
    safe_name = "".join(c for c in page_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    filename = f"{safe_name}_label.docx"
    
    # 파일 저장
    page_doc.save(filename)
    print(f"페이지 저장 완료: {filename} ({len(items_for_page)}개 라벨)")
    return True

# 상품별로 그룹화하여 처리
current_product = None
current_items = []
total_files_created = 0

while idx < len(items):
    name, price, category, code = items[idx]
    
    # 새로운 상품이 시작되거나 현재 상품의 라벨이 78개에 도달한 경우
    if current_product != name or len(current_items) >= labels_per_page:
        # 이전 상품의 라벨들이 있으면 파일로 저장
        if current_items:
            if create_label_page(current_items, current_product):
                total_files_created += 1
            current_items = []
        
        current_product = name
    
    # 현재 아이템을 리스트에 추가
    current_items.append((name, price, category, code))
    idx += 1

# 마지막 상품 처리
if current_items:
    if create_label_page(current_items, current_product):
        total_files_created += 1

print(f"\n=== 작업 완료 ===")
print(f"총 {total_files_created}개 파일 생성됨")
print(f"총 {len(items)}개 라벨 처리됨")